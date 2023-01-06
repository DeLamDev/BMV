import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
from extr_empresas import lista_empresas
import xlrd
from xlutils.copy import copy

meses = ['enero', 'febrero', 'marzo', 'abril', 
         'mayo', 'junio', 'julio', 'agosto', 
         'septiembre', 'octubre', 'noviembre', 
         'diciembre']

def fecha_actual():
    return date.today()

def fecha_elegida():
    print("¿De qué fecha desea extraer los eventos?\n")
    print("Ingrese la fecha con el formato: día-mes-año (ejem. 31-01-2000)")
    print("Si desea la fecha de hoy favor de escribir 'h'.")
    fecha_select = input("Fecha: ")
    return fecha_select
    
def validar_fecha(fecha, meses):
    format = "%d-%m-%Y"
    se_fecha = fecha.strip()
    if fecha.strip().lower() == 'h':
        print(f"Usted ingresó: {fecha_actual().strftime(format)}\n")
        return fecha_actual(), meses[int(fecha_actual().strftime('%m')) - 1]
    try:
        se_fecha = datetime.strptime(se_fecha, format)
        print(f"Usted ingresó: {fecha}\n")
        return se_fecha, meses[int(se_fecha.strftime('%m')) -1]
    except Exception as e:
        print("No ha ingresado la fecha con el formato correcto.")
        print("Recuerde ingresar el día, mes y año separados por un '-' sin espacios entre ellos.")
        print("Ejemplo: 01-10-2000 (primero de octubre del año 2000).\n")
        print(e)
        quit()

def titulo_y_sub():
    titulo = input("Favor de ingresar título del documento: ")
    subtitulo = input("Favor de ingresar subtítulo del documento: ")
    print(f"El título elegido es: {titulo}")
    print(f"El subtítulo elegido es: {subtitulo}")
    return titulo, subtitulo

def links_EvRel(perfil):
    base = "https://www.bmv.com.mx"
    r = requests.get(perfil)
    soup = BeautifulSoup(r.text, "html.parser")
    for e in soup.find_all("div", class_='tabs-area')[0].find_all('a'):
        if 'relevantevents' in e.get('href'):
            return base + e.get('href')

def extractorXls():
    book = xlrd.open_workbook("listaEmpresas.xls")
    sh = book.sheet_by_index(0)
    final_dict = {}
    for r in range(1, sh.nrows):
        final_dict[sh.cell_value(rowx=r, colx=0)] = []
        for c in range(1, sh.ncols):
            final_dict[sh.cell_value(rowx=r, colx=0)].append(sh.cell_value(rowx=r, colx=c))
    return final_dict

def relv_adder():
    b = xlrd.open_workbook("listaEmpresas.xls")
    book = copy(b)
    sh = book.get_sheet(0)
    counter = 1
    dict = extractorXls()
    for e in dict:
        sh.write(counter, 3, links_EvRel(dict[e][1]))
        counter += 1
    book.save('listaEmpresas.xls')

def doc_creator(fecha, mes, title, subtitle, dict):
    documento = Document()
    fecha_doc = documento.add_paragraph(fecha.strftime('%d') + ' de ' + mes + ' de ' + fecha.strftime('%Y'))
    fecha_doc.runs[0].bold = True
    fecha_doc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    titulo = documento.add_paragraph(title)
    if title != "":
        titulo.runs[0].bold = True
    documento.add_paragraph(subtitle)
    tabla = documento.add_table(rows=len(dict) + 1, cols=5)
    tabla.style = 'TableGrid'
    filas = tabla.rows[0].cells
    filas[0].text = "Número"
    filas[1].text = "Clave emisora"
    filas[2].text = "Razón social"
    filas[3].text = "Eventos relevantes"
    for n in range(1, len(dict) + 1):
        tabla.rows[n].cells[0].text = str(n)
        tabla.rows[n].cells[1].text = list(dict.keys())[n - 1]
        tabla.rows[n].cells[2].text = dict[list(dict.keys())[n - 1]][0]
    documento.save("Reporte_EV_BMV.docx")

def doc_updater(resultados):
    try:
        doc = Document('Reporte_EV_BMV.docx')
    except Exception as e:
        print("Es necesario generar primero el documento y tenerlo en el mismo directorio desde el que se ejecuta.")
        print(e)
        exit()
    tabla = doc.tables[0]
    doc_cleaner(tabla)
    count = 0
    for elemento in resultados:
        count += 1
        for empresa in elemento:
            formateo_eventos(tabla.cell(count, 3), elemento, empresa)
    doc.save('Reporte_EV_BMV.docx')

def doc_cleaner(tabla):
    for cell in range(1, len(tabla.rows)):
        tabla.cell(cell, 3).text = ""

def formateo_eventos(cell, elemento, empresa):
    if len(elemento[empresa]) == 3:
        texto = cell.add_paragraph()
        texto.add_run("Evento: ").bold = True
        texto.add_run(elemento[empresa][1])
        texto.add_run("\nFecha: ").bold = True
        texto.add_run(elemento[empresa][0])
        texto.add_run("\nLinks: ").bold = True
        for l in elemento[empresa][2]:
            texto.add_run("\n" + "https://www.bmv.com.mx" + l)
    elif len(elemento[empresa]) > 3:
        for n in range(0, len(elemento[empresa]) + 1, 3):
            texto = cell.add_paragraph()
            texto.add_run("Evento: ").bold = True
            texto.add_run(elemento[empresa][n - 2])
            texto.add_run("\nFecha: ").bold = True
            texto.add_run(elemento[empresa][n -3])
            texto.add_run("\nLinks: ").bold = True
            for l in elemento[empresa][n -1]:
                texto.add_run("\n" + "https://www.bmv.com.mx" + l)
            texto.add_run("\n")
    elif len(elemento[empresa]) == 0:
        cell.add_paragraph("Sin publicaciones.")
    else:
        cell.add_paragraph("Error, favor de verificar manualmente.")

def loop_events(tables, number):
    busqueda = {}
    for event in tables[number].find_all("tr"):
        for count, data in enumerate(event.find_all("td")):
            if count == 0:
                busqueda[data.text] = []
            if count == 2:
                busqueda[list(busqueda)[-1]].append([link["href"] for link in data.find_all("a")])
            elif count == 1:
                busqueda[list(busqueda)[-1]].append(data.text)
    return busqueda

def rel_event_extractor(link):
    re = requests.get(link)
    soup = BeautifulSoup(re.text, "html.parser")
    tables = soup.find_all("tbody")
    results = []
    if len(tables) > 1:
        for e in range(1, len(tables)):
            results.append(loop_events(tables, e))
        return results
    else:
        return None

def searcher(dict, fecha):
    lista_word = []
    for empresa in dict:
        lista_word.append({empresa: []})
        resultados = rel_event_extractor(dict[empresa][2])
        if resultados == None:
            continue
        for e in resultados:
            for i in e:
                if i.split()[0] == fecha.strftime("%d-%m-%Y"):
                    print(f"Evento encontrado de empresa: {empresa}.")
                    lista_word[-1][empresa].append(i)
                    lista_word[-1][empresa].extend(e[i])
    return lista_word


def checker_link_ER():
    book = xlrd.open_workbook('listaEmpresas.xls')
    sh = book.sheet_by_index(0)
    for e in range(1, sh.nrows):
        if sh.cell_value(rowx=e, colx=3) == "":
            print("Faltan links en el excel.")
            print("Favor de utilizar la función para completarlos o hacerlo de forma manual.")
            exit()

def confirmar():
    while True:
        decision = input("¿Confirmar lo ingresado? (s/n): ")
        if decision == "s":
            return True
        elif decision == "n":
            return False
        else:
            print("Favor de solo respoder con s/n.\n")
            

if __name__ == "__main__":
    print("Extractor de eventos relevantes de empresas de la BMV.\n\n")
    while True:
        elec_fecha = fecha_elegida()
        fecha, mes = validar_fecha(elec_fecha, meses)
        if confirmar():
            break
    while True:
        elec_lista = input("¿Deséa crear una lista de empresas? (s/n): ")
        if elec_lista == 's':
            lista_empresas()
            break
        elif elec_lista == 'n':
            break
        else:
            print("Favor de solo responder: s/n.\n")
    while True:
        titulo, subtitulo = titulo_y_sub()
        if confirmar():
            break
    while True:
        dec_EVREL = input("Desea añadir los links de eventos relevantes (s/n): ")
        if dec_EVREL.strip().lower() == "s":
            print("Añadiendo links, favor de esperar...")
            relv_adder()
            new_dict = extractorXls()
            print("Procesando, favor de esperar...")
            resultados = searcher(new_dict, fecha)
            break
        elif dec_EVREL.strip().lower() == "n":
            checker_link_ER()
            new_dict = extractorXls()
            print("Procesando, favor de esperar...")
            resultados = searcher(new_dict, fecha)
            break
        else:
            print("Favor de solo responder s/n.\n")
    while True:
        dec_EVREL = input("¿Generar nuevo documento? (s/n): ")
        if dec_EVREL.strip().lower() == "s":
            doc_creator(fecha, mes, titulo, subtitulo, new_dict)
            doc_updater(resultados)
            break
        elif dec_EVREL.strip().lower() == "n":
            doc_updater(resultados)
            break
        else:
            print("Favor de solo responder s/n.\n")

